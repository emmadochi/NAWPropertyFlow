import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_masterpiece_document():
    doc = Document()
    
    # Page setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B) # Slate 800

    # Header / Title Block
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    header_table.columns[0].width = Inches(6.9)
    cell = header_table.cell(0, 0)
    set_cell_background(cell, "0B132B") # Deep Luxury Navy
    set_cell_margins(cell, top=320, bottom=320, left=320, right=320)

    p_org = cell.paragraphs[0]
    p_org.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_org = p_org.add_run("NAW WORLD TECHNOLOGIES LIMITED")
    r_org.bold = True
    r_org.font.size = Pt(10)
    r_org.font.color.rgb = RGBColor(0x38, 0xBD, 0xF8) # Sky blue
    
    p_sub = cell.add_paragraph()
    r_sub = p_sub.add_run("Enterprise Software Division • NAW Property Flow CRM")
    r_sub.font.size = Pt(9.5)
    r_sub.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8) # Gray 400

    p_title = cell.add_paragraph()
    r_title = p_title.add_run("ENTERPRISE CRM PRODUCT ARCHITECTURE, OPERATIONAL MANUAL & GOVERNANCE SPECIFICATION")
    r_title.bold = True
    r_title.font.size = Pt(16.5)
    r_title.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p_for = cell.add_paragraph()
    r_for = p_for.add_run("Prepared Exclusively for the Executive Leadership & Staff of: RICAF Nigeria Limited (crm.ricafltd.com)")
    r_for.font.size = Pt(10.5)
    r_for.font.color.rgb = RGBColor(0xFE, 0xA5, 0x00) # Brand Amber/Orange
    r_for.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 1. Executive Summary
    h1 = doc.add_heading(level=1)
    r1 = h1.add_run("1. Executive Overview & Strategic Vision")
    r1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    doc.add_paragraph(
        "NAW Property Flow CRM is a high-performance, enterprise-grade real estate operating system engineered by "
        "NAW World Technologies Limited specifically to power the sales, development, financial auditing, and human capital "
        "operations of RICAF Nigeria Limited. "
        "The platform establishes an automated, end-to-end digital nervous system that eliminates operational friction, accelerates deal velocity, "
        "secures payment compliance, and delivers a world-class luxury property ownership experience to local and diaspora investors."
    )

    # 2. Business Impact & ROI Matrix
    h2 = doc.add_heading(level=1)
    r2 = h2.add_run("2. Business Transformation & ROI Matrix")
    r2.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph(
        "By replacing manual spreadsheets, scattered paper receipts, and delayed approvals with NAW Property Flow CRM, "
        "RICAF Nigeria Limited realizes measurable operational advantages across key commercial dimensions:"
    )

    roi_table = doc.add_table(rows=1, cols=3)
    roi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    roi_table.autofit = False
    widths_roi = [Inches(1.8), Inches(2.5), Inches(2.6)]
    
    headers_roi = ["Operational Dimension", "Traditional Manual Operations", "NAW Property Flow CRM Impact"]
    for i, title in enumerate(headers_roi):
        cell_h = roi_table.rows[0].cells[i]
        cell_h.width = widths_roi[i]
        set_cell_background(cell_h, "1E293B")
        set_cell_margins(cell_h, top=120, bottom=120, left=120, right=120)
        p = cell_h.paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    roi_data = [
        (
            "Marketer Lead Ownership & Anti-Theft",
            "Internal disputes, client poaching, and lack of transparency over lead origination.",
            "100% Tamper-Proof. Leads are permanently locked to the capturing sales executive upon entry/import."
        ),
        (
            "Receipt Issuance & Dispatch Time",
            "24 to 48 hours of manual Word/Excel formatting, printing, signing, and emailing.",
            "Instant (0 Seconds). Branded, stamped PDF receipts (REC-XXXXXX) auto-generate and email immediately."
        ),
        (
            "Financial Audit & Commission Integrity",
            "Overpayments, unverified commissions paid on bounced transfers, spreadsheet errors.",
            "Two-Tier Audit Safeguard: Marketer commissions strictly queue until Admin audits bank funds."
        ),
        (
            "Diaspora & Remote Investor Trust",
            "High client anxiety over unverified progress, lost paper documents, repetitive support calls.",
            "24/7 1-Tap Client Portal on mobile: Real-time milestone progress, photo feeds, and statements."
        ),
        (
            "Monthly Staff & Commission Payroll",
            "3 to 5 working days of tedious accounting, recalculating commission tiers by hand.",
            "1-Click Automated Generation: Base salary + verified commissions compiled in under 10 seconds."
        ),
        (
            "Marketing & Lead Conversion",
            "Cold outreach with no tracking, unorganized WhatsApp broadcasts, lost follow-ups.",
            "Automated Drip Sequences, Lead Scoring, Overdue Follow-Up Alerts, and Click/Open Analytics."
        )
    ]

    for dim, before, after in roi_data:
        row_cells = roi_table.add_row().cells
        for i, text in enumerate([dim, before, after]):
            row_cells[i].width = widths_roi[i]
            set_cell_background(row_cells[i], "F8FAFC" if i % 2 == 0 else "FFFFFF")
            set_cell_margins(row_cells[i], top=100, bottom=100, left=100, right=100)
            p = row_cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(text)
            r.font.size = Pt(9)
            if i == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            elif i == 2:
                r.font.color.rgb = RGBColor(0x05, 0x96, 0x69) # Emerald Green bold accent

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 3. Visual System Architecture & End-to-End Workflow Diagram
    h3 = doc.add_heading(level=1)
    r3 = h3.add_run("3. End-to-End Operational Lifecycle & Data Flow")
    r3.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph(
        "The following structured flowchart illustrates how customer prospects transition seamlessly "
        "through the CRM lifecycle into closed sales, audited funds, automated commissions, and lifelong investor portal engagement:"
    )

    diagram_text = (
        "┌────────────────────────────────────────────────────────────────────────────────────────┐\n"
        "│                             1. LEAD ACQUISITION & CAPTURE                              │\n"
        "│  • Channels: Website Inquiries, WhatsApp, Social Media Ads, Cold Calls, Bulk CSV Import│\n"
        "│  • Safeguard: Permanent ownership lock to the originating Sales Executive (Marketer)   │\n"
        "└───────────────────────────────────────────┬────────────────────────────────────────────┘\n"
        "                                            ▼\n"
        "┌────────────────────────────────────────────────────────────────────────────────────────┐\n"
        "│                           2. LEAD NURTURING & KANBAN STAGING                           │\n"
        "│  • Stages: New ➔ Contacted ➔ Follow-Up Scheduled ➔ Site Inspection Booked ➔ Negotiation│\n"
        "│  • Real-time Top KPI Counters: Total Pipeline, New Prospects, Conversion Rate %        │\n"
        "└───────────────────────────────────────────┬────────────────────────────────────────────┘\n"
        "                                            ▼\n"
        "┌────────────────────────────────────────────────────────────────────────────────────────┐\n"
        "│                         3. DEAL CLOSING & PART-PAYMENT SPREAD                          │\n"
        "│  • Unit Reserving: Specific plot selected (Available ➔ Reserved ➔ Sold)               │\n"
        "│  • Live Calculator: 100% Outright vs Part-Payment Spread (3 to 24 Months)              │\n"
        "│  • Milestone Scheduling: Milestone 1 (Deposit) receipted; remaining tranches scheduled │\n"
        "└───────────────────────────────────────────┬────────────────────────────────────────────┘\n"
        "                                            ▼\n"
        "┌────────────────────────────────────────────────────────────────────────────────────────┐\n"
        "│                  4. INSTANT PDF RECEIPT & 1-TAP CLIENT PORTAL SHARING                  │\n"
        "│  • Official Stamped PDF Receipt (REC-XXXXXX) auto-emailed to buyer inbox               │\n"
        "│  • 1-Click WhatsApp Sharing: 64-Hex Cryptographic Magic Link sent to client's phone   │\n"
        "│  • Buyer Portal: Client tracks live site photos, allocation letters, and balances 24/7 │\n"
        "└───────────────────────────────────────────┬────────────────────────────────────────────┘\n"
        "                                            ▼\n"
        "┌────────────────────────────────────────────────────────────────────────────────────────┐\n"
        "│                   5. ADMIN PAYMENT VERIFICATION & COMMISSION APPROVAL                  │\n"
        "│  • Two-Tier Financial Audit: Payment sits as 'Pending Audit' until bank credit verified│\n"
        "│  • Admin Sign-Off: Super Admin / Company Admin clicks 'Verify Payment'                │\n"
        "│  • Commission Queuing: Marketer commission moves to 'Approved' status                  │\n"
        "└───────────────────────────────────────────┬────────────────────────────────────────────┘\n"
        "                                            ▼\n"
        "┌────────────────────────────────────────────────────────────────────────────────────────┐\n"
        "│                  6. MONTHLY PAYROLL GENERATION & LEADERBOARD RECOGNITION               │\n"
        "│  • 1-Click Monthly Payroll: Aggregates Base Salaries + Approved Commissions into PDF   │\n"
        "│  • Bank Schedule Export: Instant CSV for electronic banking disbursement               │\n"
        "│  • Sales Leaderboard: Live 🥇 Gold, 🥈 Silver, 🥉 Bronze rankings & targets             │\n"
        "└────────────────────────────────────────────────────────────────────────────────────────┘"
    )

    p_diag = doc.add_paragraph()
    r_diag = p_diag.add_run(diagram_text)
    r_diag.font.name = 'Consolas'
    r_diag.font.size = Pt(8)
    r_diag.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p_diag.paragraph_format.line_spacing = 1.05
    p_diag.paragraph_format.space_after = Pt(12)

    # 4. Role-Based Security Matrix
    h4 = doc.add_heading(level=1)
    r4 = h4.add_run("4. Role-Based Security Matrix & Permissions")
    r4.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph(
        "The system enforces strict access control across all 5 operational roles to ensure data privacy and prevent unauthorized actions:"
    )

    roles_table = doc.add_table(rows=1, cols=3)
    roles_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    roles_table.autofit = False
    widths = [Inches(1.5), Inches(2.2), Inches(3.2)]
    
    headers = ["Role Identifier", "Access Scope", "Core Responsibilities & Safeguards"]
    for i, title in enumerate(headers):
        cell_h = roles_table.rows[0].cells[i]
        cell_h.width = widths[i]
        set_cell_background(cell_h, "1E293B")
        set_cell_margins(cell_h, top=120, bottom=120, left=120, right=120)
        p = cell_h.paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    roles_data = [
        (
            "Super Admin & Company Admin",
            "Company-Wide & Multi-Branch",
            "• Absolute visibility over all company revenues, financial charts, and audit logs.\n"
            "• Sole authorized sign-off on Payment Verifications before commissions disburse.\n"
            "• 1-Click Monthly Payroll approval and automated bank payment schedule export.\n"
            "• System settings, branch creation, user permissions, and legal template configuration."
        ),
        (
            "Sales Manager",
            "Branch & Team Leadership",
            "• Real-time visibility across branch leads on visual Kanban & Table pipelines.\n"
            "• Managerial authority to assign and reallocate leads among active sales agents.\n"
            "• Schedule and approve site inspections and track branch monthly sales quotas.\n"
            "• Monitor agent conversion metrics and lead follow-up compliance."
        ),
        (
            "Sales Executive (Marketer)",
            "Personal Pipeline & Personal Earnings",
            "• Guaranteed Lead Ownership: All leads created or imported via CSV are locked to their user ID.\n"
            "• Data Privacy: Only their assigned prospects are visible; other agents' contacts and company payroll financials are hidden.\n"
            "• Real-time commission earnings tracking on their personalized dashboard card.\n"
            "• 1-Click WhatsApp Client Portal sharing with buyers."
        ),
        (
            "HR & Payroll Manager",
            "Staff Governance, Appraisals & Payroll",
            "• Track company-wide staff leaderboard rankings, sales volumes, and revenue.\n"
            "• Automated Monthly Payroll Run: Aggregates base salaries + approved marketer commissions.\n"
            "• Manage employee leave applications, staff submissions, certifications, and onboarding."
        ),
        (
            "Buyer / Investor (Client Portal)",
            "Personal Property Portfolio",
            "• 1-Tap Magic Link Access: Instant zero-password login from WhatsApp or Email.\n"
            "• Financial Statement: Real-time tracking of Amount Paid vs Outstanding Balance.\n"
            "• Download stamped PDF receipts (REC-XXXXXX) and Allocation/Contract documents.\n"
            "• Real-time inspection of estate development milestones and progress photos."
        )
    ]

    for role_name, scope, desc in roles_data:
        row_cells = roles_table.add_row().cells
        for i, text in enumerate([role_name, scope, desc]):
            row_cells[i].width = widths[i]
            set_cell_background(row_cells[i], "F8FAFC" if i % 2 == 0 else "FFFFFF")
            set_cell_margins(row_cells[i], top=100, bottom=100, left=100, right=100)
            p = row_cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(text)
            r.font.size = Pt(9)
            if i == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 5. Exhaustive Feature Directory (13 Core Modules)
    h5 = doc.add_heading(level=1)
    r5 = h5.add_run("5. Exhaustive Core Feature Directory")
    r5.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    modules = [
        (
            "1. Lead Capture, Auto-Assignment & Drag-and-Drop Pipeline",
            "• Single & Bulk CSV Ingestion: Quick creation and spreadsheet upload with sample templates.\n"
            "• Locked Marketer Ownership: Automatic assignment protection preventing unauthorized reallocation.\n"
            "• Dual-View Pipeline: Seamless switching between Table and visual Drag-and-Drop Kanban boards.\n"
            "• Executive Top Counters: Real-time overview of Total Pipeline, New Prospects, Tours in Flight, and Conversion Rate %."
        ),
        (
            "2. Estate, Plot & Unit Inventory Engine",
            "• Master Estates & Projects: Categorize properties by Estate Name, Property Type (Land, Duplex, Terrace), Price, and Location.\n"
            "• Unit-Level Inventory: Track specific plot and unit numbers (e.g. Plot 12, Block B).\n"
            "• Interactive Unit Lifecycle: Unit status transitions seamlessly between 'Available' -> 'Reserved' -> 'Sold / Converted'.\n"
            "• Bulk Unit Generator: 1-click batch generation of plots/units across entire phases."
        ),
        (
            "3. Inspections & Follow-Up Compliance Tracking",
            "• Booking Calendar: Schedule physical or virtual site inspections with clients.\n"
            "• Officer Assignment: Assign inspection chaperones and log feedback notes.\n"
            "• Structured Follow-Ups: Schedule Call, WhatsApp, Email, or Meeting follow-ups with automatic overdue reminders.\n"
            "• Timeline Audit: Every client touchpoint is permanently recorded on their activity log."
        ),
        (
            "4. Dynamic Part-Payment, Installment Calculator & Milestone Engine",
            "• Flexible Payment Structures: Toggle between 100% Outright and Part-Payment plans.\n"
            "• Live Installment Calculator: Instant calculation of Deposit Paid, Spread Duration (3–24 Months), Remaining Balance, and Monthly Tranches.\n"
            "• Automated Milestone Scheduling: Milestone 1 (Deposit) is receipted immediately while future tranches automatically schedule reminder due dates."
        ),
        (
            "5. Instant Stamped PDF Receipt Generation & Email Dispatch",
            "• Automated Document Generation: Immediate generation of official PDF receipts (REC-XXXXXX) upon payment entry.\n"
            "• Corporate Stamping: Features official RICAF verification stamps, bank transaction reference, amount paid in words, outstanding balance, and estate allocation.\n"
            "• Automated Outgoing Email Delivery: Directly dispatches stamped PDF receipts to the investor's inbox via authenticated SMTP."
        ),
        (
            "6. Payment Verification Safeguards & Automated Commission Queuing",
            "• Two-Tier Financial Audit: Payments enter as 'Pending Audit' until confirmed by Admin.\n"
            "• Admin Sign-Off: Only Company Admins or Super Admins can verify payments.\n"
            "• Automated Marketer Commission: Recalculates verified funds, records verifier timestamp, and queues commission with status 'Approved' for monthly payroll."
        ),
        (
            "7. 1-Tap Client Portal Magic Link & WhatsApp Sharing",
            "• 0-Friction Mobile Access: Marketers share a 1-tap WhatsApp or Email link directly from the client profile.\n"
            "• 256-Bit Cryptographic Token: 64-hex random security token provides instant password-free access.\n"
            "• Live Client Portfolio: Investors view payment statements, download receipts, and inspect estate construction progress 24/7."
        ),
        (
            "8. Legal Document Generation & Contract Automation",
            "• Dynamic Variable Templates: Build Deed of Contract, Provisional Allocation Letters, and Receipts using merge tags like {{buyer_name}}, {{property_name}}, and {{amount_paid}}.\n"
            "• 1-Click PDF Generation: Instant compiling of legal contracts ready for digital download and email delivery."
        ),
        (
            "9. Targeted Newsletter Campaigns & Marketing Drip Sequences",
            "• Audience Segmentation: Filter recipients by Lead Status (e.g. Hot Prospects) or Property Interest.\n"
            "• Rich Visual Builder: Clean HTML newsletters with real-time recipient counts.\n"
            "• Automated Drip Sequences: Trigger automated email sequences based on prospect lifecycle stages.\n"
            "• Open & Click Analytics: Full tracking on email engagement."
        ),
        (
            "10. Dynamic Multi-Department Management & Custom KPI Governance (Zero-Code Scalability)",
            "• Infinite Organizational Scalability: Create and manage unlimited operational departments (e.g. Media & Creative, Administration, Legal & Land Titles, Customer Relations, Logistics, Security) without writing new code.\n"
            "• Custom KPI & Metric Configurator: Define unique performance metrics for each department (e.g., 'Site Videos Shot', 'Engagement Numbers', 'Deed of Assignment Drafted', 'Office Maintenance Turnaround %').\n"
            "• Hierarchical Staff Submissions: Staff submit daily/weekly performance logs against their departmental metrics and attach supporting files.\n"
            "• HOD Approval & MD Scorecards: Head of Departments review and approve submissions, which instantly feed into the Executive Consolidated Department Performance Report."
        ),
        (
            "11. HR Governance, Sales Leaderboard & 1-Click Automated Payroll",
            "• Gamified Sales Leaderboard: Live monthly rankings highlighting Gold, Silver, and Bronze sales champions, revenue totals, and conversion rates.\n"
            "• Departmental Target Tracking: Set revenue and unit quotas per department.\n"
            "• 1-Click Automated Monthly Payroll: Automatically aggregates base staff salaries + all approved monthly marketer commissions into consolidated payslips ready for bank export.\n"
            "• Staff Governance: Employee leave applications, staff onboarding checklists, and disciplinary/review logs."
        ),
        (
            "12. Cloud File Storage & Digital Asset Vault",
            "• Enterprise Document Repository: Create structured folders for Estate Layouts, Survey Plans, Land Titles, and Corporate Documents.\n"
            "• Granular Actions: Upload, preview, rename, and download high-resolution architectural files securely."
        ),
        (
            "13. Virtual 3D Tour Integration (Prototype)",
            "• Interactive Visual Walkthrough: Embedded virtual tour exploration for prospective buyers to experience estate environments remotely."
        ),
        (
            "14. Multi-Branch Operations & Consolidated Reporting",
            "• Multi-Branch Switching: Filter the entire CRM across Head Office, Island Branch, Mainland Branch, or regional offices.\n"
            "• Executive Export Suite: 1-click CSV/Excel exports for Leads by Source, Sales by Agent, Follow-Up Compliance, and Branch Comparisons."
        )
    ]

    for m_title, m_desc in modules:
        h_m = doc.add_heading(level=2)
        r_m = h_m.add_run(m_title)
        r_m.font.size = Pt(11.5)
        r_m.font.color.rgb = RGBColor(0xFE, 0xA5, 0x00)
        
        p = doc.add_paragraph(m_desc)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(8)

    # 6. Step-by-Step SOP Workflows
    h6 = doc.add_heading(level=1)
    r6 = h6.add_run("6. Standard Operating Procedures (SOP)")
    r6.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    sops = [
        ("SOP 1: Inbound Lead Capture & Bulk CSV Import", 
         "1. Navigate to 'Leads' -> Click 'Add Lead Prospect' or 'Import Leads (CSV)'.\n"
         "2. Fill in buyer name, phone, WhatsApp number, budget range, and preferred estate.\n"
         "3. If logged in as a Sales Executive, ownership is automatically locked to you.\n"
         "4. Click 'Create Lead' to populate the prospect in the pipeline."),
        
        ("SOP 2: Reserving a Unit & Recording a Deal with Part-Payment Spread",
         "1. Open the Lead profile (/leads/{id}) -> Click 'Record Sale'.\n"
         "2. Choose the Property Unit and select Payment Structure ('Part-Payment / Spread').\n"
         "3. Enter 'Deposit Paid Today' (e.g. ₦3,000,000) and 'Spread Duration' (e.g. 6 Months).\n"
         "4. Enter the Bank Teller / Transfer Reference and click 'Close Deal & Generate Plan'.\n"
         "5. The system issues Milestone 1 as Paid, generates the official stamped PDF receipt, and emails it to the buyer."),

        ("SOP 3: Admin Payment Verification & Marketer Commission Approval",
         "1. Admin visits 'Sales & Payments' -> 'Milestones'.\n"
         "2. Locate the payment marked 'Pending Audit' -> Verify bank credit.\n"
         "3. Click 'Verify Payment (Admin)' -> Confirm the prompt.\n"
         "4. The payment is stamped 'Verified' and marketer commission is approved for monthly payroll."),

        ("SOP 4: Sharing the 1-Tap Client Portal with the Buyer",
         "1. Open the Lead profile (/leads/{id}).\n"
         "2. Locate the 'Client Portal Access' card on the left panel.\n"
         "3. Click 'Share on WhatsApp' to send a pre-formatted invitation link to the buyer's phone, or click 'Copy Access Link' to email it."),

        ("SOP 5: Generating Legal Allocation & Contract Documents",
         "1. Open 'Legal & Documents' -> 'Generated Documents' -> Click 'Generate Document'.\n"
         "2. Select the Lead and the Template (e.g. Deed of Agreement / Allocation Letter).\n"
         "3. The system compiles the dynamic PDF, populating client details, plot allocation, and payment history.\n"
         "4. 1-Click 'Email to Client' or download the signed document."),

        ("SOP 6: Running Monthly Staff Payroll & Exporting Bank Schedules",
         "1. HR / Admin visits 'HR Management' -> 'Payroll'.\n"
         "2. Click 'Generate Monthly Payroll' -> Select Month & Year.\n"
         "3. The system automatically computes Base Salary + All Verified Commissions.\n"
         "4. Click 'Download Payslips' to print or email official PDF payslips to all staff, and export the bank transfer CSV.")
    ]

    for s_title, s_desc in sops:
        h_s = doc.add_heading(level=2)
        r_s = h_s.add_run(s_title)
        r_s.font.size = Pt(11)
        r_s.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        p = doc.add_paragraph(s_desc)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)

    # 7. Troubleshooting, FAQs & Field Exception Protocols
    h7 = doc.add_heading(level=1)
    r7 = h7.add_run("7. Troubleshooting & Field Exception Protocols")
    r7.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    faqs = [
        ("Q1: What happens if a buyer makes a payment with an unclear transfer description?",
         "Resolution Protocol: The marketer logs the payment with the available bank transaction reference and notes. The milestone enters 'Pending Audit'. The Company Admin or Finance Officer cross-checks with the corporate bank statement before clicking 'Verify Payment'. Commissions are held safely until positive bank confirmation."),

        ("Q2: Can a Sales Executive change who a lead is assigned to?",
         "Resolution Protocol: No. To prevent lead tampering or internal commission disputes, sales executives cannot modify the 'assigned_to' field. If an agent goes on leave or leaves the company, only a Sales Manager or Company Admin can reassign the lead."),

        ("Q3: What if a client requests to extend their part-payment spread (e.g., from 6 to 12 months)?",
         "Resolution Protocol: An Admin can navigate to 'Payments & Milestones' -> 'Payment Plan' -> Adjust the remaining milestone tranches to spread across the newly agreed timeline."),

        ("Q4: How does a diaspora client access their portal if they lose their WhatsApp chat?",
         "Resolution Protocol: The marketer or admin can open the client's file in the CRM (/leads/{id}) and click 'Copy Access Link' or 'Share on WhatsApp' again. The 64-hex cryptographic token regenerates seamlessly without resetting any client data.")
    ]

    for q, a in faqs:
        h_q = doc.add_heading(level=2)
        r_q = h_q.add_run(q)
        r_q.font.size = Pt(10.5)
        r_q.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        p = doc.add_paragraph(a)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)

    # 8. Enterprise Security, Privacy & Data Compliance (NDPR)
    h8 = doc.add_heading(level=1)
    r8 = h8.add_run("8. Enterprise Security, Data Privacy & NDPR Compliance")
    r8.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph(
        "NAW Property Flow CRM is built in strict adherence to modern cybersecurity standards and the "
        "Nigeria Data Protection Regulation (NDPR):"
    )

    sec_points = (
        "• Cryptographic Token Security: Client portal magic links use 256-bit entropy random tokens (bin2hex(random_bytes(32))), mathematically impervious to brute-force guessing.\n"
        "• Transport Layer Security (TLS/SSL): 100% of data traffic, client documents, and login credentials are encrypted in transit via 256-bit SSL.\n"
        "• Multi-Tenant Data Isolation: Database queries are tenant-scoped, ensuring complete isolation of customer data from any external systems.\n"
        "• Automated Server Backups & Disaster Recovery: Automated database snapshots run continuously on cPanel server infrastructure.\n"
        "• Comprehensive Audit Trail: Critical operational events (logins, payment records, role changes, document downloads) are timestamped and logged."
    )
    p_sec = doc.add_paragraph(sec_points)
    p_sec.paragraph_format.line_spacing = 1.2
    p_sec.paragraph_format.space_after = Pt(10)

    # 9. User Onboarding & Fast-Track Implementation Roadmap
    h9 = doc.add_heading(level=1)
    r9 = h9.add_run("9. 30-Day Onboarding & Implementation Roadmap")
    r9.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    roadmap_table = doc.add_table(rows=1, cols=3)
    roadmap_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    roadmap_table.autofit = False
    widths_rm = [Inches(1.3), Inches(2.2), Inches(3.4)]
    
    headers_rm = ["Phase / Timeline", "Focus Area", "Milestones & Key Deliverables"]
    for i, title in enumerate(headers_rm):
        cell_h = roadmap_table.rows[0].cells[i]
        cell_h.width = widths_rm[i]
        set_cell_background(cell_h, "1E293B")
        set_cell_margins(cell_h, top=120, bottom=120, left=120, right=120)
        p = cell_h.paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    rm_data = [
        ("Week 1 (Days 1–7)", "System Setup & Inventory Population", "Create user accounts (Admins, Managers, Sales Execs), configure branch offices, and upload master estates & plot units."),
        ("Week 2 (Days 8–14)", "Sales Executive Field Training", "Hands-on staff workshop on lead entry, bulk CSV imports, Kanban pipeline management, and 1-tap WhatsApp portal sharing."),
        ("Week 3 (Days 15–21)", "Finance, Audit & Legal Verification", "Finance team dry-run on recording part-payments, generating stamped PDF receipts, admin payment verification, and contract generation."),
        ("Week 4 (Days 22–30)", "Go-Live, Leaderboards & First Payroll Run", "Full operational go-live across all branches, launching live leaderboard rankings, and executing the first automated monthly payroll.")
    ]

    for phase, focus, deliver in rm_data:
        row_cells = roadmap_table.add_row().cells
        for i, text in enumerate([phase, focus, deliver]):
            row_cells[i].width = widths_rm[i]
            set_cell_background(row_cells[i], "F8FAFC" if i % 2 == 0 else "FFFFFF")
            set_cell_margins(row_cells[i], top=100, bottom=100, left=100, right=100)
            p = row_cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(text)
            r.font.size = Pt(9)
            if i == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 10. Technical Support Escalation Hierarchy & SLA
    h10 = doc.add_heading(level=1)
    r10 = h10.add_run("10. Technical Support Hierarchy & Service Level Agreement (SLA)")
    r10.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph(
        "NAW World Technologies Limited provides multi-tiered technical support to ensure 99.9% uptime for RICAF Nigeria Limited:"
    )

    sla_data = (
        "• Tier-1 (User Support & Onboarding Guidance): Response within 2 hours. Guidance on feature navigation, password resets, and user roles.\n"
        "• Tier-2 (Application & Workflow Configuration): Response within 4 hours. Template customizations, department target updates, and report exports.\n"
        "• Tier-3 (Critical Infrastructure & Security Escalations): Immediate dispatch (under 30 minutes). Server performance, SSL certificate renewals, database failover, and SMTP routing.\n\n"
        "Official Support Channels:\n"
        "📧 Enterprise Support Email: info@nawtechnologies.com\n"
        "🌐 Production Server Instance: https://crm.ricafltd.com"
    )
    p_sla = doc.add_paragraph(sla_data)
    p_sla.paragraph_format.line_spacing = 1.2
    p_sla.paragraph_format.space_after = Pt(12)

    # 11. Official Handover & Acceptance Sign-Off
    h11 = doc.add_heading(level=1)
    r11 = h11.add_run("11. Product Acceptance & Formal Handover Sign-Off")
    r11.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph(
        "This document certifies that the NAW Property Flow CRM application has been developed, configured, tested, "
        "and successfully deployed on the live production server (crm.ricafltd.com) according to the operational specifications of RICAF Nigeria Limited."
    )

    sign_table = doc.add_table(rows=1, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_table.autofit = False
    widths_s = [Inches(3.4), Inches(3.4)]

    s_headers = ["FOR: NAW WORLD TECHNOLOGIES LIMITED\n(Technology Provider & System Architect)", "FOR: RICAF NIGERIA LIMITED\n(Client & Operating Company)"]
    for i, title in enumerate(s_headers):
        cell_h = sign_table.rows[0].cells[i]
        cell_h.width = widths_s[i]
        set_cell_background(cell_h, "1E293B")
        set_cell_margins(cell_h, top=140, bottom=140, left=140, right=140)
        p = cell_h.paragraphs[0]
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    sign_row = sign_table.add_row().cells
    sign_text_naw = (
        "\nAuthorized Signatory: ___________________________\n\n"
        "Name: Lead Software Architect & Director\n\n"
        "Designation: Enterprise Systems Division\n\n"
        "Date: August 20, 2026\n\n"
        "Official Seal / Stamp:\n\n\n"
    )
    sign_text_ricaf = (
        "\nAuthorized Signatory: ___________________________\n\n"
        "Name: Managing Director / COO\n\n"
        "Designation: Executive Management\n\n"
        "Date: _________________________\n\n"
        "Official Seal / Stamp:\n\n\n"
    )

    for i, text in enumerate([sign_text_naw, sign_text_ricaf]):
        sign_row[i].width = widths_s[i]
        set_cell_background(sign_row[i], "FAFAFA")
        set_cell_margins(sign_row[i], top=120, bottom=120, left=120, right=120)
        p = sign_row[i].paragraphs[0]
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Footer Table
    foot_table = doc.add_table(rows=1, cols=1)
    foot_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    foot_table.columns[0].width = Inches(6.9)
    fcell = foot_table.cell(0, 0)
    set_cell_background(fcell, "F1F5F9")
    set_cell_margins(fcell, top=140, bottom=140, left=140, right=140)
    pf = fcell.paragraphs[0]
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run("DOCUMENT PREPARED BY NAW WORLD TECHNOLOGIES LIMITED • NAW PROPERTY FLOW CRM DIVISION\nClient Support & Deployments: info@nawtechnologies.com • Live Server Instance: crm.ricafltd.com")
    rf.font.size = Pt(8.5)
    rf.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Save to file
    out_dir = r"c:\xampp\htdocs\NAWPropertyFlowCRM"
    out_path = os.path.join(out_dir, "NAW_Property_Flow_CRM_RICAF_Product_Guide.docx")
    doc.save(out_path)
    print(f"Successfully generated master enterprise manual: {out_path}")

if __name__ == '__main__':
    create_masterpiece_document()
